"""
Apply manuscript formatting to a pandoc-generated docx:
  - Times New Roman 12 pt throughout
  - 1.5 line spacing
  - Justified alignment
  - All black text
  - No paragraph spacing (before/after = 0)
  - Page numbers in footer (centred)
  - Explicit run-level bold on specific section headings (preserved across edits)
  - Page break before specific section headings (manuscript "front matter -> Main text")
"""
import sys
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
SIZE = Pt(11)

# Headings whose runs get explicit bold added on top of the Heading 2 style.
# Match: heading text (case-insensitive, stripped) STARTS WITH one of these prefixes —
# permissive so e.g. "Abstract (~150 words; unstructured)" still bolds via the "abstract" prefix.
HEADINGS_TO_BOLD_PREFIX = ("main text", "acknowledgements", "abstract", "authors")

# Headings that should start on a new page (page break BEFORE the heading) — same prefix logic.
HEADINGS_WITH_PAGEBREAK_PREFIX = ("main text", "abstract")


def set_run_format(run):
    run.font.name = FONT
    run.font.size = SIZE
    run.font.color.rgb = BLACK
    # Force east-asian / complex-script font (some Word installs ignore .name otherwise)
    rPr = run._element.get_or_add_rPr()
    for tag in ("rFonts",):
        old = rPr.find(qn("w:" + tag))
        if old is not None:
            rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rPr.append(rFonts)


def set_paragraph_format(p):
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    for run in p.runs:
        set_run_format(run)


def insert_page_break_before(p):
    """Insert a page break as the first run of paragraph p."""
    run = p.add_run()
    # Move the run we just added to the front
    p._element.insert(list(p._element).index(run._element), run._element)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)
    set_run_format(run)


def apply_heading_overrides(doc):
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if "eading" not in style.lower():
            continue
        key = p.text.strip().lower()
        if not key:
            continue
        if any(key.startswith(prefix) for prefix in HEADINGS_TO_BOLD_PREFIX):
            for run in p.runs:
                run.bold = True
        if any(key.startswith(prefix) for prefix in HEADINGS_WITH_PAGEBREAK_PREFIX):
            already = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        already = True
            if not already:
                br_run = OxmlElement("w:r")
                br = OxmlElement("w:br")
                br.set(qn("w:type"), "page")
                br_run.append(br)
                p._element.insert(0, br_run)


def walk_paragraphs(doc):
    # Body paragraphs
    for p in doc.paragraphs:
        set_paragraph_format(p)
    # Paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_format(p)
    apply_heading_overrides(doc)


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        # Replace existing footer paragraphs with a single centred PAGE field
        for p in list(footer.paragraphs):
            p.clear()
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        # Field code for PAGE
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        set_run_format(run)


def main(in_path, out_path):
    doc = Document(in_path)

    # Update Normal style as a baseline
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    rPr = normal.element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rPr.append(rFonts)

    # Update heading styles to also be Times New Roman + black, kept bold for hierarchy
    for h in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"):
        try:
            st = doc.styles[h]
            st.font.name = FONT
            st.font.color.rgb = BLACK
            if h == "Title":
                st.font.size = Pt(13)
            elif h == "Heading 1":
                st.font.size = Pt(12)
            else:
                st.font.size = Pt(11)
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after = Pt(0)
            st.paragraph_format.line_spacing = 1.5
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            rPr = st.element.get_or_add_rPr()
            for old in rPr.findall(qn("w:rFonts")):
                rPr.remove(old)
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), FONT)
            rFonts.set(qn("w:hAnsi"), FONT)
            rFonts.set(qn("w:cs"), FONT)
            rFonts.set(qn("w:eastAsia"), FONT)
            rPr.append(rFonts)
        except KeyError:
            pass

    walk_paragraphs(doc)
    add_page_number_footer(doc)

    doc.save(out_path)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
